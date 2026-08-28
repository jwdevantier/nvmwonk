#!/usr/bin/env python3
# SPDX-FileCopyrightText: 2025 Jesper Wendel Devantier <jwd@defmacro.it>
#
# SPDX-License-Identifier: BSD-2-Clause

"""
nvmwonk — parse and query NVMe / NVM specification documents.

Subcommands:
  extract [options] INPUT.docx OUTPUT.xml
      Parse a .docx (Technical Proposal or published spec) into structured
      XML. See `nvmwonk extract --help` for the schema.

  query figures FILES...
      List every figure (captioned tables + graphic captions) across files.

  query figure REF FILES...
      Extract one figure's <table>, verbatim, by number ("31") or token
      ("FIGCDW11"); falls back to substring match on the caption.

  query section REF FILES...
      Extract a section (number like '3.1.3', '5.NEW.1', 'Annex A', or
      heading substring) with all its content (paragraphs, tables,
      subsections) verbatim.

  query xpath 'EXPR' FILES...
      Raw XPath 1.0 escape hatch (EXSLT re: namespace available).

All query commands emit one well-formed XML document on stdout;
diagnostics go to stderr. Exit 0 when anything matched, 1 on no match,
2 on usage/XPath error. Every result carries a <breadcrumb> of its
ancestor sections so the caller knows where in the spec a hit lives.

Schema cheatsheet (for crafting XPath expressions):
  tp > metadata > title
  tp > body > (section | spec-changes | annex)*  — recursive <section>
  section@heading, spec-changes@spec, annex@heading
  p                                              — ordinary paragraphs
  table@caption > row(@header="true")? > cell    — verbatim spec tables
  inline markup (TPs only): content@added="true", tbd, note, br
"""

import argparse
import os
import re
import sys
from dataclasses import dataclass
from typing import Optional
from lxml import etree
import docx
from docx.oxml.ns import qn
from docx.text.run import Run

_RUN_CONTAINERS = ("fldSimple", "hyperlink", "ins", "smartTag")


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  Extract module                                                       ║
# ╚══════════════════════════════════════════════════════════════════════╝



def iter_runs(para):
    """Yield docx Run objects for `para` in document order, including runs
    nested inside w:fldSimple / w:hyperlink / w:ins / w:smartTag."""
    for child in para._p:
        tag = etree.QName(child).localname
        if tag == "r":
            yield Run(child, para)
        elif tag in _RUN_CONTAINERS:
            for r in child.iter(qn("w:r")):
                yield Run(r, para)


# ── List-number reconstruction ─────────────────────────────────────────────
#
# The published specs number their headings via style-level multilevel list
# numbering (word/numbering.xml), which python-docx never renders: NVMe H1..H6
# → numId 154 ilvl 0..5 (lvlText %1, %1.%2, ...), Annex/Annex.1/Annex.1.1 →
# numId 109 ilvl 0..2 (lvlText 'Annex %1.' upperLetter, '%1.%2', '%1.%2.%3.').
# NumberingState replays those counters while walking the document so the XML
# headings carry the numbers readers actually see (e.g. "3.1.3 Controller
# Types", "Annex A. Sanitize Operation Considerations (Informative)").

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _roman(n: int) -> str:
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
            (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
            (5, "V"), (4, "IV"), (1, "I")]
    out = []
    for v, s in vals:
        while n >= v:
            out.append(s)
            n -= v
    return "".join(out)


def _fmt_counter(n: int, fmt: str) -> str:
    if fmt == "upperLetter":
        s = ""
        while n > 0:
            n, r = divmod(n - 1, 26)
            s = chr(65 + r) + s
        return s
    if fmt == "lowerLetter":
        return _fmt_counter(n, "upperLetter").lower()
    if fmt == "upperRoman":
        return _roman(n)
    if fmt == "lowerRoman":
        return _roman(n).lower()
    return str(n)  # decimal and anything unexpected


class NumberingState:
    """Replays style-driven multilevel list numbering for heading styles.

    Build once per document via build(); then call format_for(style_name) for
    each heading paragraph to get the number to prepend (or None when the
    style carries no numbering — e.g. TP headings, whose numbers are literal
    text already).
    """

    def __init__(self, style_map, num_map):
        # style_map: style name -> (numId, ilvl)
        # num_map:   numId -> {ilvl: {"start": int, "fmt": str, "text": str}}
        self.style_map = style_map
        self.num_map = num_map
        self.counters = {}  # numId -> [int] per ilvl (0..8)

    @staticmethod
    def build(doc) -> "NumberingState":
        # numbering.xml part
        numbering_part = None
        for rel in doc.part.rels.values():
            if "numbering" in rel.reltype:
                numbering_part = rel.target_part
                break
        if numbering_part is None:
            return NumberingState({}, {})
        root = etree.fromstring(numbering_part.blob)

        num_to_abstract = {}
        for num in root.findall(f"{{{_W}}}num"):
            an = num.find(f"{{{_W}}}abstractNumId")
            if an is not None:
                num_to_abstract[num.get(f"{{{_W}}}numId")] = an.get(f"{{{_W}}}val")

        abstract_levels = {}   # abstractId -> {ilvl: {...}}
        pstyle_link = {}       # styleId -> (abstractId, ilvl)
        for an in root.findall(f"{{{_W}}}abstractNum"):
            aid = an.get(f"{{{_W}}}abstractNumId")
            levels = {}
            for lvl in an.findall(f"{{{_W}}}lvl"):
                ilvl = int(lvl.get(f"{{{_W}}}ilvl"))
                def gv(tag):
                    e = lvl.find(f"{{{_W}}}{tag}")
                    return e.get(f"{{{_W}}}val") if e is not None else None
                levels[ilvl] = {
                    "start": int(gv("start") or 1),
                    "fmt": gv("numFmt") or "decimal",
                    "text": gv("lvlText") or f"%{ilvl + 1}",
                }
                ps = gv("pStyle")
                if ps:
                    pstyle_link[ps] = (aid, ilvl)
            abstract_levels[aid] = levels

        # style name -> (numId, ilvl), resolved through basedOn and the
        # pStyle links in numbering.xml.
        style_ids = {}  # styleId -> style element
        for s in doc.styles:
            sid = s.element.get(qn("w:styleId"))
            if sid:
                style_ids[sid] = s.element

        def style_numid(el, depth=0):
            """Effective numId for a style, walking basedOn; None if absent."""
            if el is None or depth > 10:
                return None
            ppr = el.find(qn("w:pPr"))
            numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
            if numpr is not None:
                nid = numpr.find(qn("w:numId"))
                if nid is not None and nid.get(qn("w:val")) not in (None, "0"):
                    return nid.get(qn("w:val"))
            bo = el.find(qn("w:basedOn"))
            if bo is not None:
                return style_numid(style_ids.get(bo.get(qn("w:val"))), depth + 1)
            return None

        heading_styles = set(STYLE_DEPTH) | set(ANNEX_STYLE_DEPTH)
        style_map = {}
        for s in doc.styles:
            if s.name not in heading_styles:
                continue
            el = s.element
            sid = el.get(qn("w:styleId"))
            numid = style_numid(el)
            if numid is None or numid not in num_to_abstract:
                continue
            aid = num_to_abstract[numid]
            # Prefer the abstractNum's own pStyle -> ilvl link (authoritative);
            # fall back to the style's explicit w:ilvl, else 0.
            if sid in pstyle_link and pstyle_link[sid][0] == aid:
                ilvl = pstyle_link[sid][1]
            else:
                ppr = el.find(qn("w:pPr"))
                numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
                il = numpr.find(qn("w:ilvl")) if numpr is not None else None
                ilvl = int(il.get(qn("w:val"))) if il is not None else 0
            style_map[s.name] = (numid, ilvl)

        # NVMe H7 resolves (via basedOn Heading7) to a numId whose ilvl-0
        # would render as a bare counter — wrong. Real H7 headings carry
        # explicit paragraph-level numPr into the NVMe H* list and are
        # handled by the list-item path in heading_info(); drop the bogus
        # style mapping so any numPr-less H7 paragraph simply gets no number.
        if "NVMe H7" in style_map and "NVMe H1" in style_map:
            if style_map["NVMe H7"][0] != style_map["NVMe H1"][0]:
                del style_map["NVMe H7"]

        num_map = {numid: abstract_levels[aid]
                   for numid, aid in num_to_abstract.items()}
        return NumberingState(style_map, num_map)

    def heading_numids(self) -> set:
        """numIds that drive heading numbering in this document."""
        return {numid for numid, _ in self.style_map.values()}

    def format_explicit(self, numid: str, ilvl: int) -> Optional[str]:
        """Advance counters for an explicitly-numbered paragraph (numPr on
        the paragraph itself, e.g. the base spec's NVMe H7 headings)."""
        if numid not in self.num_map:
            return None
        return self._advance(numid, ilvl)

    def format_for(self, style_name: str) -> Optional[str]:
        """Advance counters for a heading of this style; return the rendered
        number (e.g. '3.1.3' or 'Annex A.'), or None if unnumbered."""
        if style_name not in self.style_map:
            return None
        numid, ilvl = self.style_map[style_name]
        return self._advance(numid, ilvl)

    def _advance(self, numid: str, ilvl: int) -> Optional[str]:
        levels = self.num_map.get(numid)
        if not levels or ilvl not in levels:
            return None
        if numid not in self.counters:
            # counters indexed by ilvl; start−1 so the first increment yields
            # the level's declared start value
            self.counters[numid] = [
                levels.get(i, {"start": 1})["start"] - 1
                for i in range(max(levels) + 1)
            ]
        ctr = self.counters[numid]
        ctr[ilvl] += 1
        for deeper in range(ilvl + 1, len(ctr)):
            lv = levels.get(deeper)
            ctr[deeper] = (lv["start"] - 1) if lv else 0
        text = levels[ilvl]["text"]
        def repl(m):
            n = int(m.group(1)) - 1
            lv = levels.get(n, {"fmt": "decimal"})
            return _fmt_counter(ctr[n], lv["fmt"])
        return re.sub(r"%([1-9])", repl, text)


# ── Run classification (NVMe TP markup conventions) ────────────────────────

BLUE = "0000FF"
RED = "FF0000"
PURPLE = "7030A0"
GREEN = "00B050"


def classify_run(run) -> tuple[str, str | None]:
    """Classify a docx run by its semantic marker (TP markup convention)."""
    color = None
    if run.font.color and run.font.color.rgb:
        color = str(run.font.color.rgb)
    highlight = run.font.highlight_color is not None
    strike = run.font.strike

    # Deleted text (omit)
    if color == RED and strike:
        return ("deleted", None)
    if color == PURPLE and strike:
        return ("deleted", None)

    # TBD values (highlighted)
    if color == BLUE and highlight:
        return ("tbd", None)

    # New or moved text
    if color == BLUE:
        return ("added", None)
    if color == PURPLE:
        return ("added", None)

    # Editor notes
    if color == GREEN:
        return ("note", None)

    return ("plain", None)


# ── Run merging ──────────────────────────────────────────────────────────────

@dataclass
class TextSpan:
    marker: str
    text: str
    extra: Optional[str] = None


def merge_runs(runs) -> list[TextSpan]:
    """Merge consecutive runs with the same semantic marker, omitting deleted."""
    spans = []
    for run in runs:
        marker, extra = classify_run(run)
        if marker == "deleted":
            continue
        text = run.text
        if not text:
            continue
        if spans and spans[-1].marker == marker and spans[-1].extra == extra:
            spans[-1].text += text
        else:
            spans.append(TextSpan(marker=marker, text=text, extra=extra))
    return spans


def _append_text(parent: etree.Element, text: str):
    """Append character data to `parent` at the current position: the
    element's .text when it has no children yet, otherwise the last child's
    .tail. (Appending to .text unconditionally would serialize the text
    BEFORE existing children such as <br/> or nested <table>, silently
    reordering — and in the old cell code, losing later paragraphs.)"""
    if len(parent):
        last = parent[-1]
        last.tail = (last.tail or "") + text
    else:
        parent.text = (parent.text or "") + text


def spans_to_xml(spans: list[TextSpan], parent: etree.Element):
    """Append text/spans into a parent XML element."""
    for span in spans:
        if span.marker == "plain":
            _append_text(parent, span.text)
        elif span.marker == "added":
            el = etree.SubElement(parent, "content", added="true")
            el.text = span.text
        elif span.marker == "tbd":
            el = etree.SubElement(parent, "tbd")
            el.text = span.text
        elif span.marker == "note":
            el = etree.SubElement(parent, "note")
            el.text = span.text


def runs_to_xml(runs, parent: etree.Element):
    spans = merge_runs(runs)
    spans_to_xml(spans, parent)


def paragraph_text_plain(para) -> str:
    parts = []
    for span in merge_runs(iter_runs(para)):
        parts.append(span.text)
    return "".join(parts)


def cell_text_plain(cell) -> str:
    """Get plain text of a cell (all paragraphs)."""
    parts = []
    for para in cell.paragraphs:
        parts.append(paragraph_text_plain(para))
    return " | ".join(parts)


# ── Heading / caption / spec-changes detection ─────────────────────────────

# Word style → heading depth
STYLE_DEPTH = {
    "NVMe H1": 1, "NVMe H2": 2, "NVMe H3": 3,
    "NVMe H4": 4, "NVMe H5": 5, "NVMe H6": 6, "NVMe H7": 7,
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
}

ANNEX_STYLE_DEPTH = {
    "Annex": 1,
    "Annex.1": 2,
    "Annex.1.1": 3,
}

# Styles that are pure navigation artifacts — never body content. The base
# spec's table of figures lists every caption as "Figure N: <title><tab><page>";
# without this skip those 800+ lines land in the body as fake caption
# paragraphs and poison the figure list.
SKIP_STYLES = {"table of figures", "Table of Contents", "TOC"}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def is_list_item(para) -> bool:
    """True if the paragraph is part of a numbered/bulleted list (has w:numPr)."""
    return para._p.find(f".//{{{W_NS}}}numPr") is not None


def match_spec_changes(text: str) -> Optional[str]:
    """Return the spec name if `text` is a TP spec-changes header, else None."""
    m = re.match(r'Description of Specification Changes for the (.+?)\s*$', text)
    if m:
        return m.group(1).strip()
    return None


def paragraph_numpr(para) -> Optional[tuple[str, int]]:
    """Paragraph-level list numbering (numId, ilvl), or None."""
    numpr = para._p.find(f".//{{{W_NS}}}numPr")
    if numpr is None:
        return None
    nid = numpr.find(qn("w:numId"))
    il = numpr.find(qn("w:ilvl"))
    numid = nid.get(qn("w:val")) if nid is not None else None
    ilvl = int(il.get(qn("w:val"))) if il is not None else 0
    if not numid or numid == "0":
        return None
    return (numid, ilvl)


def heading_info(para, in_body: bool = True,
                 heading_numids: Optional[set] = None
                 ) -> Optional[tuple[int, str, str]]:
    """
    Detect whether a paragraph is a heading.

    Returns (depth, text, kind) where kind is "section" or "annex".

    Word-style detection (NVMe H1..H7, Heading 1..6, Annex*) always applies.
    The text-based fallback (e.g. "3 NVM Express Architecture") only applies
    when `in_body` is True — otherwise Normal-styled prose in the front
    matter (e.g. "3855 SW 153rd Drive") would be misclassified as a heading.

    List items (w:numPr present) styled with Word heading styles are
    ignored — those are source-document styling accidents (e.g. base-spec
    list items tagged with NVMe H1 by mistake). List items that match a
    text-based heading pattern are still recognised, because that's how
    TPs style their section headings.
    """
    style = para.style.name if para.style else ""
    text = paragraph_text_plain(para).strip()
    if not text:
        return None

    # Annex styles are trusted even on list items — an annex is rarely
    # confusable with prose content (the base spec uses "Annex" style for
    # the annex header and "Annex.1" / "Annex.1.1" for sub-annexes; these
    # titles never look like content).
    if style in ANNEX_STYLE_DEPTH:
        return (ANNEX_STYLE_DEPTH[style], text, "annex")

    # List items (w:numPr present) styled with NVMe H*/Heading N* are
    # usually source-document styling accidents (e.g. base-spec list items
    # tagged with NVMe H1 by mistake: their numPr points at numId 0 or an
    # unrelated list). Exception: paragraphs whose explicit numPr references
    # the document's *heading* multilevel list (e.g. the base spec's NVMe H7
    # headings, which all carry numPr numId=154 ilvl=6) are real headings;
    # their depth comes from the paragraph's own ilvl.
    if style in STYLE_DEPTH:
        if para._p.find(f".//{{{W_NS}}}numPr") is not None:
            numpr = paragraph_numpr(para)
            if (numpr is not None and heading_numids
                    and numpr[0] in heading_numids):
                return (numpr[1] + 1, text, "section")
            # numPr present but not part of the heading numbering scheme.
            # Two cases: (a) styling accidents — prose paragraphs tagged
            # with a heading style and a real list (base spec's NVMe H1
            # list items, numId 0/256) — the text fallback below matches
            # nothing and they stay prose; (b) TP headings whose author
            # disabled list numbering and typed the number literally
            # ('5.NEW.1 Hard Limit Mode', NVMe H3 + numId 0) — the text
            # fallback recognises the number pattern.
            pass  # fall through to the text-based fallback
        else:
            return (STYLE_DEPTH[style], text, "section")

    if not in_body:
        return None

    # Text-based fallback (TPs without Word heading styles). Use \s+ because
    # the source docs separate the section number from the title with a
    # non-breaking space (U+00A0), not a regular ASCII space. The title must
    # start with an uppercase letter: prose and formula paragraphs can begin
    # with a number too (NVM CS "2 ^ DataExponent", "16 namespace
    # granularity descriptors ..."), while real NVMe section titles are
    # Title Case ("2 NVM Command Set Model", "5.NEW.1 Hard Limit Mode").
    if re.match(r"^\d+\.\d+\.\d+\.\d+\s+[A-Z]", text):
        return (4, text, "section")
    if re.match(r"^\d+\.\d+\.\d+\s+[A-Z]", text):
        return (3, text, "section")
    if re.match(r"^\d+\.\d+\s+[A-Z]", text):
        return (2, text, "section")
    if re.match(r"^\d+\s+[A-Z]", text):
        return (1, text, "section")
    # TP "NEW" section numbers: '5.NEW' (depth 2), '5.NEW.1' (depth 3),
    # '5.NEW.1.2' (depth 4).
    m = re.match(r"^\d+\.NEW((?:\.\d+)*)\s+\S", text)
    if m:
        return (2 + m.group(1).count("."), text, "section")
    if re.match(r"^\d+\.NEW", text):
        return (2, text, "section")

    return None


def is_figure_caption(para) -> Optional[str]:
    """
    Detect whether a paragraph is a figure/table caption. Returns the caption
    text, or None.

    A caption starts with "Figure" or "Table" followed by a number/token
    and a colon (e.g. "Figure 31: …", "Figure : …", "Figure TOKENBKT: …").
    The colon must come immediately after the figure token — prose like
    "Figure 11 shows the relationship …, which has:" is not a caption.
    """
    style = para.style.name if para.style else ""
    text = paragraph_text_plain(para).strip()
    if not text:
        return None
    if style == "NVMe Figure Title":
        return text
    if re.match(r"^(Figure|Table)\s*\S*\s*:", text):
        return text
    return None


# ── Merged cell detection ────────────────────────────────────────────────────

def detect_merged_cells(row) -> list[int]:
    """
    Detect horizontally merged cells in a row.
    Returns a list of col indices that are continuations of a merge
    (i.e., should be skipped).
    """
    cells = row.cells
    skip = []
    for i in range(1, len(cells)):
        if cell_text_plain(cells[i]) == cell_text_plain(cells[i - 1]):
            skip.append(i)
    return skip


def write_table_rows(table, tbl_el, start_row: int = 0):
    """Write rows/cells of a docx table into <table> element tbl_el.

    Cells are processed in true document order at the w:tc child level, so
    nested tables (w:tbl inside w:tc — the spec's bitfield subtables inside
    register-layout figures) are preserved as nested <table> elements
    instead of being silently dropped (cell.paragraphs alone skips them).
    Recursion depth is naturally bounded by the source document.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for row_idx in range(start_row, len(table.rows)):
        row = table.rows[row_idx]
        row_el = etree.SubElement(tbl_el, "row")
        if row_idx == start_row:
            row_el.set("header", "true")

        merged_skip = detect_merged_cells(row)
        for col_idx, cell in enumerate(row.cells):
            if col_idx in merged_skip:
                continue
            cell_el = etree.SubElement(row_el, "cell")
            wrote_text = False
            for child in cell._tc:
                tag = etree.QName(child).localname
                if tag == "p":
                    cp = Paragraph(child, cell)
                    if wrote_text:
                        etree.SubElement(cell_el, "br")
                    runs_to_xml(iter_runs(cp), cell_el)
                    wrote_text = True
                elif tag == "tbl":
                    nested_el = etree.SubElement(cell_el, "table")
                    write_table_rows(Table(child, cell), nested_el)


def _current_parent(current_annex, current_spec_changes, section_stack):
    """Pick the most specific active parent element.

    The section stack's bottom is always the current container (body,
    spec-changes, or annex) and headings push onto it — so the stack top IS
    the right parent. (Formerly this returned current_spec_changes /
    current_annex directly, which flattened all TP content onto the
    spec-changes element and left every section an empty shell.)
    """
    return section_stack[-1][0]


# ── Main extraction ──────────────────────────────────────────────────────────

def extract(docx_path: str, xml_path: str, doc_id: Optional[str] = None,
            doc_title: Optional[str] = None):
    doc = docx.Document(docx_path)
    numbering = NumberingState.build(doc)

    root = etree.Element("tp")
    if doc_id:
        root.set("id", doc_id)
    if doc_title:
        root.set("title", doc_title)

    # ── Build ordered element list ──
    body_docx = doc.element.body
    elements = []
    para_idx = 0
    table_idx = 0
    for child in body_docx:
        tag = child.tag
        if tag == qn("w:p"):
            if para_idx < len(doc.paragraphs):
                elements.append(("para", para_idx))
                para_idx += 1
        elif tag == qn("w:tbl"):
            if table_idx < len(doc.tables):
                elements.append(("table", table_idx))
                table_idx += 1

    # ── Metadata: <title> (from CLI arg, or auto-extracted) ──
    meta = etree.SubElement(root, "metadata")
    title = doc_title
    if title is None:
        title_parts = []
        for p in doc.paragraphs[:50]:
            if p.style and p.style.name in ("Document Title", "Title"):
                t = paragraph_text_plain(p).strip()
                if t:
                    title_parts.append(t)
            elif title_parts:
                # Stop once we've passed the title block
                break
        if title_parts:
            title = " ".join(title_parts)
    if title:
        etree.SubElement(meta, "title").text = title

    # ── Body ──
    body_el = etree.SubElement(root, "body")

    # The "current parent" is the top of section_stack. The stack holds
    # (element, depth) pairs. When we encounter a new heading, we pop entries
    # whose depth is >= the new heading's depth, then append the new section.
    section_stack = [(body_el, 0)]
    current_spec_changes = None  # None or a <spec-changes> element (TPs only)
    current_annex = None         # None or an <annex> element

    # Front matter: skip content until we hit a spec-changes header or any
    # Word-styled heading. This naturally skips title pages, mailing
    # addresses, legal notices, the "Markup Conventions" section, etc.
    in_body = False

    i = 0
    while i < len(elements):
        kind, idx = elements[i]

        if kind == "para":
            para = doc.paragraphs[idx]
            if para.style and para.style.name in SKIP_STYLES:
                i += 1
                continue
            text = paragraph_text_plain(para).strip()
            if not text:
                i += 1
                continue

            # TP spec-changes header — always triggers the start of body
            spec_name = match_spec_changes(text)
            if spec_name:
                in_body = True
                current_annex = None
                current_spec_changes = etree.SubElement(
                    body_el, "spec-changes", spec=spec_name
                )
                section_stack = [(current_spec_changes, 0)]
                i += 1
                continue

            # Heading? Word-style headings always count; text-based fallback
            # only applies once we're already inside the body.
            info = heading_info(para, in_body=in_body,
                                heading_numids=numbering.heading_numids())
            if info:
                in_body = True
                depth, heading_text, kind_h = info

                # Prepend the reconstructed list number (published specs
                # number headings via style-level numbering, which
                # python-docx never renders). Skipped for TP headings -
                # their styles carry no numbering, and their numbers are
                # literal text already (double-numbering guard). The guard
                # requires a *complete* number (digits, optional dotted
                # groups, then whitespace, or a .NEW section): a title may
                # itself start with a digit (NVM CS "16b Guard Protection
                # Information") and must still get its number prepended.
                numpr = paragraph_numpr(para)
                if numpr is not None and numpr[0] in numbering.heading_numids():
                    num = numbering.format_explicit(*numpr)
                else:
                    num = numbering.format_for(
                        para.style.name if para.style else "")
                if num and not re.match(r"^(\d+(\.\d+)*\s|\d+\.NEW|Annex\s+[A-Z])",
                                      heading_text):
                    heading_text = f"{num} {heading_text}"

                if kind_h == "annex":
                    if depth == 1:
                        # Top-level annex — wrap in <annex>
                        current_spec_changes = None
                        current_annex = etree.SubElement(
                            body_el, "annex", heading=heading_text
                        )
                        section_stack = [(current_annex, 0)]
                    else:
                        # Sub-annex (Annex.1, Annex.1.1) — nest as <section>
                        while len(section_stack) > 1 and section_stack[-1][1] >= depth:
                            section_stack.pop()
                        new_section = etree.SubElement(
                            section_stack[-1][0], "section", heading=heading_text
                        )
                        section_stack.append((new_section, depth))
                else:
                    # Regular section heading
                    if current_annex is not None and depth <= 1:
                        # A new top-level section after the annex(es) — close
                        # the annex and reset to body
                        current_annex = None
                        section_stack = [(body_el, 0)]
                    while len(section_stack) > 1 and section_stack[-1][1] >= depth:
                        section_stack.pop()
                    new_section = etree.SubElement(
                        section_stack[-1][0], "section", heading=heading_text
                    )
                    section_stack.append((new_section, depth))
                i += 1
                continue

            # Figure caption? Also marks the start of body content.
            cap = is_figure_caption(para)
            if cap:
                in_body = True
                # Emit the caption as a paragraph. We deliberately do NOT carry
                # it forward as "last_caption" for the next table — a figure
                # caption often describes a graphic (not a table), and tables
                # carry their own caption in the merged first row.
                p_el = etree.SubElement(_current_parent(
                    current_annex, current_spec_changes, section_stack
                ), "p")
                runs_to_xml(iter_runs(para), p_el)
                i += 1
                continue

            # Regular paragraph
            if not in_body:
                i += 1
                continue
            p_el = etree.SubElement(_current_parent(
                current_annex, current_spec_changes, section_stack
            ), "p")
            runs_to_xml(iter_runs(para), p_el)
            i += 1

        elif kind == "table":
            if not in_body:
                i += 1
                continue
            table = doc.tables[idx]
            tbl_el = etree.SubElement(_current_parent(
                current_annex, current_spec_changes, section_stack
            ), "table")

            # Determine caption. Two sources, in priority order:
            #   1. First row is a single merged cell starting with "Figure"
            #      or "Table" (this row IS the caption row and will be
            #      skipped).
            #   2. The immediately preceding element is a paragraph whose
            #      text is a Figure/Table caption.
            caption = None
            skip_first_row = False
            first_row_cells = table.rows[0].cells
            first_row_unique = set(cell_text_plain(c) for c in first_row_cells)
            # The caption row is usually one merged cell spanning the whole
            # table, but some docs pad it with empty cells (NVM CS Figure
            # 123: caption + one stray empty cell; TP Figure 105: two empty
            # cells + the caption). Treat the row as a caption row whenever
            # its non-empty cells collapse to a single text starting with
            # "Figure"/"Table".
            nonempty = {t for t in first_row_unique if t.strip()}
            if len(nonempty) == 1:
                cap_text = next(iter(nonempty))
                if cap_text and re.match(r"^(Figure|Table)\b", cap_text):
                    caption = cap_text
                    skip_first_row = True
            if caption is None and i > 0 and elements[i - 1][0] == "para":
                prev_para = doc.paragraphs[elements[i - 1][1]]
                prev_text = paragraph_text_plain(prev_para).strip()
                if is_figure_caption(prev_para) is not None:
                    caption = prev_text
            if caption is not None:
                tbl_el.set("caption", caption)

            start_row = 1 if skip_first_row else 0
            write_table_rows(table, tbl_el, start_row)
            i += 1

    # ── Write XML ──
    tree = etree.ElementTree(root)
    tree.write(xml_path, pretty_print=True, xml_declaration=True, encoding="UTF-8")
    print(f"Wrote {xml_path}")




# ╔══════════════════════════════════════════════════════════════════════╗
# ║  Query module                                                         ║
# ╚══════════════════════════════════════════════════════════════════════╝


NSMAP = {"re": "http://exslt.org/regular-expressions"}

MAX_XPATH_RESULTS = 200

CAPTION_RE = re.compile(r"^(Figure|Table)\s*(\S*)\s*:\s*(.*)$")


CAPTION_RE = re.compile(r"^(Figure|Table)\s*(\S*)\s*:\s*(.*)$")


def parse_caption(caption: str):
    """Split a caption into (kind, num, title); num may be a token or ''."""
    m = CAPTION_RE.match(caption.strip())
    if not m:
        return (None, None, caption.strip())
    return (m.group(1), m.group(2), m.group(3))


def breadcrumb(node) -> etree.Element:
    """<breadcrumb> with nested ancestor context elements, outermost first."""
    chain = []
    cur = node.getparent()
    while cur is not None:
        tag = etree.QName(cur).localname
        if tag == "section":
            chain.append(("section", "heading", cur.get("heading")))
        elif tag == "annex":
            chain.append(("annex", "heading", cur.get("heading")))
        elif tag == "spec-changes":
            chain.append(("spec-changes", "spec", cur.get("spec")))
        cur = cur.getparent()
    bc = etree.Element("breadcrumb")
    for tag, attr, val in reversed(chain):
        etree.SubElement(bc, tag, **{attr: val or ""})
    return bc


def result_element(node, docname: str) -> etree.Element:
    """Wrap a deep copy of `node` in a <result> with doc + breadcrumb."""
    import copy
    res = etree.Element("result", doc=docname)
    res.append(breadcrumb(node))
    node_copy = copy.deepcopy(node)
    node_copy.tail = None  # keep pretty-printing clean inside <result>
    res.append(node_copy)
    return res


def load(path: str) -> etree._ElementTree:
    return etree.parse(path)


# ── figures ────────────────────────────────────────────────────────────────

def cmd_figures(paths) -> int:
    root = etree.Element("figures")
    count = 0
    for path in paths:
        docname = os.path.basename(path)
        tree = load(path)
        # 1. captioned tables
        for tbl in tree.getroot().iter("table"):
            caption = tbl.get("caption")
            if not caption:
                continue
            kind, num, _title = parse_caption(caption)
            fig = etree.SubElement(
                root, "figure", doc=docname, kind="table",
                caption=caption)
            if num:
                fig.set("num", num)
            fig.append(breadcrumb(tbl))
            count += 1
        # 2. graphic captions (caption-classified paragraphs)
        for p in tree.getroot().iter("p"):
            text = "".join(p.itertext()).strip()
            if not CAPTION_RE.match(text):
                continue
            kind, num, _title = parse_caption(text)
            fig = etree.SubElement(
                root, "figure", doc=docname, kind="paragraph", caption=text)
            if num:
                fig.set("num", num)
            fig.append(breadcrumb(p))
            count += 1
    root.set("count", str(count))
    write_out(root)
    return 0 if count else 1


# ── figure <ref> ───────────────────────────────────────────────────────────

def cmd_figure(ref, paths) -> int:
    root = etree.Element("results", ref=ref)
    count = 0
    for path in paths:
        docname = os.path.basename(path)
        tree = load(path)
        tables = [t for t in tree.getroot().iter("table") if t.get("caption")]
        # Pass 1: exact num match ("31" == "31", "FIGCDW11" == "FIGCDW11").
        exact = [t for t in tables
                 if parse_caption(t.get("caption"))[1] == ref]
        # Pass 2: substring fallback on the full caption.
        matches = exact or [t for t in tables
                            if ref.lower() in t.get("caption").lower()]
        for tbl in matches:
            res = result_element(tbl, docname)
            _kind, num, _title = parse_caption(tbl.get("caption"))
            if num:
                res.set("num", num)
            root.append(res)
            count += 1
        # Pass 3: graphic-only figures (caption paragraphs, no table).
        # The image content isn't in the XML — return the caption paragraph
        # plus the prose around it (the describing text almost always sits
        # immediately before/after the caption), and the breadcrumb so the
        # caller can pull the whole section if needed.
        if not matches:
            paras = []
            for p in tree.getroot().iter("p"):
                text = "".join(p.itertext()).strip()
                _kind, num, _title = parse_caption(text)
                if num is not None:
                    paras.append((p, text, num))
            exact_p = [(p, text, num) for p, text, num in paras if num == ref]
            hit_p = exact_p or [(p, text, num) for p, text, num in paras
                                if ref.lower() in text.lower()]
            for p, text, num in hit_p:
                res = result_element(p, docname)
                res.set("kind", "paragraph")
                if num:
                    res.set("num", num)
                # surrounding prose: up to 2 preceding + 1 following
                # sibling paragraphs (skip the image-only empties — they
                # have no text by construction)
                ctx = etree.Element("context")
                sibs = list(p.getparent())
                i = sibs.index(p)
                before = [s for s in sibs[:i] if s.tag == "p"][-2:]
                after = [s for s in sibs[i + 1:] if s.tag == "p"][:1]
                import copy
                for s in before + after:
                    c = copy.deepcopy(s)
                    c.tail = None
                    ctx.append(c)
                if len(ctx):
                    res.append(ctx)
                root.append(res)
                count += 1
    root.set("count", str(count))
    write_out(root)
    return 0 if count else 1


# ── section <ref> ──────────────────────────────────────────────────────────

def _num_prefix_match(ref: str, heading: str) -> bool:
    """True when `ref` matches the leading number/annex token of `heading`,
    delimited by whitespace or end-of-string — so '3.1' matches
    '3.1 NVM Controller Architecture' but NOT '3.1.3 Controller Types'."""
    return heading == ref or heading.startswith(ref + " ") or \
        heading.startswith(ref + "\u00a0")


def cmd_section(ref, paths) -> int:
    root = etree.Element("results", ref=ref)
    count = 0
    for path in paths:
        docname = os.path.basename(path)
        tree = load(path)
        containers = [e for e in tree.getroot().iter("section", "annex")
                      if e.get("heading")]
        # Pass 1: leading-number match ('3.1.3', '5.NEW.1', 'Annex A').
        exact = [e for e in containers
                 if _num_prefix_match(ref, e.get("heading"))]
        # Pass 2: substring fallback on the whole heading.
        matches = exact or [e for e in containers
                            if ref.lower() in e.get("heading").lower()]
        for el in matches:
            root.append(result_element(el, docname))
            count += 1
    root.set("count", str(count))
    write_out(root)
    return 0 if count else 1


# ── xpath ──────────────────────────────────────────────────────────────────

def cmd_xpath(expr, paths) -> int:
    root = etree.Element("results", xpath=expr)
    count = 0
    truncated = False
    for path in paths:
        docname = os.path.basename(path)
        tree = load(path)
        try:
            hits = tree.xpath(expr, namespaces=NSMAP)
        except etree.XPathError as e:
            print(f"XPath error: {e}", file=sys.stderr)
            return 2
        # XPath functions like count() / string() / boolean() / number()
        # return a single value, not a list — normalise to a 1-element list.
        if not isinstance(hits, list):
            hits = [hits]
        for hit in hits:
            if count >= MAX_XPATH_RESULTS:
                truncated = True
                break
            if isinstance(hit, etree._Element):
                root.append(result_element(hit, docname))
                count += 1
            elif isinstance(hit, etree._ElementUnicodeResult):
                # attribute or text node: attribute name is available for
                # attributes via .attrname / .getparent()
                parent = hit.getparent()
                if getattr(hit, "attrname", None):
                    res = etree.Element("result", doc=docname)
                    if parent is not None:
                        res.append(breadcrumb(parent))
                    etree.SubElement(res, "attribute",
                                     name=hit.attrname, value=str(hit))
                    root.append(res)
                else:
                    res = etree.Element("result", doc=docname)
                    if parent is not None:
                        res.append(breadcrumb(parent))
                    t = etree.SubElement(res, "text")
                    t.text = str(hit)
                    root.append(res)
                count += 1
            else:
                # numbers / booleans
                res = etree.SubElement(root, "result", doc=docname)
                v = etree.SubElement(res, "value")
                v.text = str(hit)
                count += 1
        if truncated:
            break
    root.set("count", str(count))
    if truncated:
        root.append(etree.Element("truncated",
                                  max=str(MAX_XPATH_RESULTS)))
    write_out(root)
    return 0 if count else 1


def write_out(root):
    try:
        sys.stdout.write(etree.tostring(root, pretty_print=True,
                                        xml_declaration=True,
                                        encoding="UTF-8").decode())
        sys.stdout.flush()
    except BrokenPipeError:
        # consumer (e.g. `head`) closed the pipe — exit quietly
        import os as _os
        _os._exit(0)



# ╔══════════════════════════════════════════════════════════════════════╗
# ║  CLI                                                                  ║
# ╚══════════════════════════════════════════════════════════════════════╝


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="nvmwonk",
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = p.add_subparsers(dest="cmd", required=True, metavar="COMMAND")

    # extract ────────────────────────────────────────────────────────────
    pe = sub.add_parser(
        "extract",
        help="Parse a .docx into structured XML.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    pe.add_argument("input", help="Input .docx file")
    pe.add_argument("output", help="Output .xml file")
    pe.add_argument("-i", "--id",
                    help='Document ID (e.g. "TP4176", "NVMe-Base-2.3")')
    pe.add_argument("-t", "--title", help="Document title")
    pe.set_defaults(func=_run_extract)

    # query ──────────────────────────────────────────────────────────────
    pq = sub.add_parser(
        "query",
        help="Query an extracted XML file.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    qsub = pq.add_subparsers(dest="query_cmd", required=True, metavar="QUERY_CMD")

    pqf = qsub.add_parser("figures",
                          help="List all figures across files")
    pqf.add_argument("files", nargs="+", help="Extracted .xml files")
    pqf.set_defaults(func=_run_query_figures)

    pqr = qsub.add_parser("figure",
                          help="Extract one figure verbatim by ref")
    pqr.add_argument("ref",
                     help='Figure number ("31"), token ("FIGCDW11"), '
                          'or caption substring')
    pqr.add_argument("files", nargs="+", help="Extracted .xml files")
    pqr.set_defaults(func=_run_query_figure)

    pqs = qsub.add_parser("section",
                          help="Extract a section verbatim by ref")
    pqs.add_argument("ref",
                     help='Section number ("3.1.3", "Annex A") or '
                          'heading substring')
    pqs.add_argument("files", nargs="+", help="Extracted .xml files")
    pqs.set_defaults(func=_run_query_section)

    pqx = qsub.add_parser("xpath",
                          help="Run an XPath 1.0 expression")
    pqx.add_argument("expr", help='XPath 1.0 (EXSLT "re:" namespace available)')
    pqx.add_argument("files", nargs="+", help="Extracted .xml files")
    pqx.set_defaults(func=_run_query_xpath)

    return p


def _run_extract(args) -> int:
    extract(args.input, args.output, doc_id=args.id, doc_title=args.title)
    return 0


def _run_query_figures(args) -> int:
    return cmd_figures(args.files)


def _run_query_figure(args) -> int:
    return cmd_figure(args.ref, args.files)


def _run_query_section(args) -> int:
    return cmd_section(args.ref, args.files)


def _run_query_xpath(args) -> int:
    return cmd_xpath(args.expr, args.files)


def main(argv=None) -> int:
    args = build_parser().parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    sys.exit(main())
